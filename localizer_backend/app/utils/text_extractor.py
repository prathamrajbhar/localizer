"""
Text Extraction Utility
Extracts text content from various file formats (PDF, DOCX, TXT)
NO FALLBACKS - Real text extraction only
"""
import os
from typing import Optional, Dict, Any
from pathlib import Path

# Import text extraction libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.utils.logger import app_logger


class TextExtractor:
    """Extract text content from various file formats"""
    
    def __init__(self):
        app_logger.info("Text Extractor initialized")
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are available"""
        missing_deps = []
        
        if not PDF_AVAILABLE:
            missing_deps.append("PyPDF2")
        if not PDFPLUMBER_AVAILABLE:
            missing_deps.append("pdfplumber")
        if not DOCX_AVAILABLE:
            missing_deps.append("python-docx")
        
        if missing_deps:
            app_logger.warning(f"Missing text extraction dependencies: {missing_deps}")
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from file based on extension
        
        Args:
            file_path: Path to file
        
        Returns:
            Dict with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.txt':
                return self._extract_from_txt(file_path)
            elif file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            app_logger.error(f"Text extraction failed for {file_path}: {e}")
            raise RuntimeError(f"Failed to extract text from {file_path.name}: {str(e)}")
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first, then fallback to other encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    app_logger.info(f"TXT extracted with {encoding} encoding: {len(text)} chars")
                    
                    return {
                        "text": text.strip(),
                        "word_count": len(text.split()),
                        "char_count": len(text),
                        "encoding": encoding,
                        "pages": 1,
                        "format": "txt"
                    }
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode text file with any common encoding")
            
        except Exception as e:
            raise RuntimeError(f"TXT extraction error: {str(e)}")
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file using multiple methods"""
        
        # Try pdfplumber first (better for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                return self._extract_pdf_pdfplumber(file_path)
            except Exception as e:
                app_logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        if PDF_AVAILABLE:
            try:
                return self._extract_pdf_pypdf2(file_path)
            except Exception as e:
                app_logger.error(f"PyPDF2 extraction also failed: {e}")
        
        raise RuntimeError("No PDF extraction libraries available or all failed")
    
    def _extract_pdf_pdfplumber(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using pdfplumber (better for complex layouts)"""
        import pdfplumber
        
        text_parts = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = "\n".join(text_parts)
        
        app_logger.info(f"PDF extracted via pdfplumber: {page_count} pages, {len(full_text)} chars")
        
        return {
            "text": full_text.strip(),
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
            "pages": page_count,
            "format": "pdf",
            "method": "pdfplumber"
        }
    
    def _extract_pdf_pypdf2(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using PyPDF2 (simpler, works for basic PDFs)"""
        import PyPDF2
        
        text_parts = []
        page_count = 0
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = "\n".join(text_parts)
        
        app_logger.info(f"PDF extracted via PyPDF2: {page_count} pages, {len(full_text)} chars")
        
        return {
            "text": full_text.strip(),
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
            "pages": page_count,
            "format": "pdf",
            "method": "PyPDF2"
        }
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available. Install with: pip install python-docx")
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            
            app_logger.info(f"DOCX extracted: {len(doc.paragraphs)} paragraphs, {len(full_text)} chars")
            
            return {
                "text": full_text.strip(),
                "word_count": len(full_text.split()),
                "char_count": len(full_text),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "format": "docx"
            }
            
        except Exception as e:
            raise RuntimeError(f"DOCX extraction error: {str(e)}")
    
    def validate_file_format(self, file_path: str) -> bool:
        """
        Validate if file format is supported
        
        Args:
            file_path: Path to file
        
        Returns:
            True if supported, False otherwise
        """
        supported_formats = {'.txt', '.pdf', '.docx', '.doc'}
        file_ext = Path(file_path).suffix.lower()
        return file_ext in supported_formats
    
    def get_supported_formats(self) -> Dict[str, bool]:
        """Get supported formats and their availability"""
        return {
            "txt": True,  # Always available
            "pdf": PDF_AVAILABLE or PDFPLUMBER_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "doc": DOCX_AVAILABLE  # python-docx can handle .doc files too
        }


# Global text extractor instance
text_extractor = TextExtractor()